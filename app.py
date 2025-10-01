import streamlit as st
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import io
import json
import time

# --- Gemini API Setup ---
# The API key is set to an empty string to allow the Canvas environment to inject it at runtime.
GEMINI_API_KEY = ""
GEMINI_MODEL_NAME = "gemini-2.5-flash-preview-05-20"
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL_NAME}:generateContent?key={GEMINI_API_KEY}"

# Define sections of the project management plan
plan_sections = [
    "Executive Summary",
    "Outcomes & Vision",
    "Objectives",
    "Scope & Deliverables",
    "Governance & Team",
    "Risk Management"
]

# --- Initialize Session State ---
if "section_2_1" not in st.session_state:
    st.session_state.section_2_1 = ""
if "sections" not in st.session_state:
    st.session_state.sections = {section: "" for section in plan_sections}
if "ai_status" not in st.session_state:
    st.session_state.ai_status = ""
# New state variables for sequential flow
if "current_section_index" not in st.session_state:
    st.session_state.current_section_index = 0
if "current_section_generated" not in st.session_state:
    st.session_state.current_section_generated = False


# Note: This function uses the non-standard '__fetch__' wrapper required for some environments.
@st.cache_data(show_spinner=False)
def call_gemini_api(user_prompt, system_prompt):
    """
    Calls the Gemini API with exponential backoff for reliable generation.
    """
    payload = {
        "contents": [{"parts": [{"text": user_prompt}]}],
        "systemInstruction": {"parts": [{"text": system_prompt}]}
    }

    max_retries = 5
    delay = 1

    for attempt in range(max_retries):
        try:
            # We use the native fetch API available in the Canvas environment
            response = st.runtime.scriptrunner.add_script_run_ctx(
                lambda: st.runtime.scriptrunner.add_script_run_ctx(
                    lambda: __fetch__(GEMINI_API_URL, {
                        "method": 'POST',
                        "headers": {'Content-Type': 'application/json'},
                        "body": json.dumps(payload)
                    })
                )()
            )()
            
            # Check for API errors
            if response.status_code == 429:
                raise Exception("Rate limit exceeded.")
            if not response.ok:
                error_detail = response.json().get("error", {}).get("message", "Unknown API error")
                raise Exception(f"API Error ({response.status_code}): {error_detail}")

            result = response.json()
            candidate = result.get('candidates', [{}])[0]
            
            if candidate and candidate.get('content', {}).get('parts', [{}])[0].get('text'):
                return candidate['content']['parts'][0]['text']
            
            return "Failed to generate content. Please check the API response structure."

        except Exception as e:
            if attempt < max_retries - 1:
                print(f"Attempt {attempt + 1} failed, retrying in {delay}s: {e}")
                time.sleep(delay)
                delay *= 2
            else:
                # Log error to console, but return a user-friendly message
                print(f"Failed to call Gemini API after {max_retries} attempts. Error: {e}")
                return "AI suggestion failed due to API error. Please enter content manually."

    return "AI suggestion failed to complete."


def generate_ai_suggestion(section_title, description):
    """
    Generates content for a plan section using the Gemini API.
    """
    system_prompt = (
        "You are an expert project manager. Your task is to write a concise, professional, "
        f"and detailed section for a Project Management Plan focusing on the '{section_title}' section. "
        "The output must be pure text, ready to be dropped into a document, and should strictly "
        "relate to the project description provided."
    )
    user_prompt = (
        f"Write the content for the '{section_title}' section of a Project Management Plan. "
        f"Base the content on the following project description:\n\n---\n\n{description}"
    )
    
    # Call the API with the specific prompts
    content = call_gemini_api(user_prompt, system_prompt)
    return content

# Callback function to generate content for the current section
def run_current_ai_suggestion():
    if not st.session_state.section_2_1:
        st.session_state.ai_status = "Please enter a Project Description first (Step 1)."
        return
        
    current_section = plan_sections[st.session_state.current_section_index]
    st.session_state.ai_status = f"Generating content for: {current_section}..."
    
    with st.spinner(f"Generating content for: {current_section}..."):
        ai_content = generate_ai_suggestion(current_section, st.session_state.section_2_1)
        st.session_state.sections[current_section] = ai_content
        st.session_state.current_section_generated = True
        st.session_state.ai_status = f"Suggestion loaded for {current_section}."
        
# Callback function to accept content and move to the next section
def accept_and_move_next():
    # Only move forward if we are not on the last section
    if st.session_state.current_section_index < len(plan_sections) - 1:
        st.session_state.current_section_index += 1
        st.session_state.current_section_generated = False
        st.session_state.ai_status = ""
    else:
        # We are on the last section, mark flow as complete
        st.session_state.current_section_index = len(plan_sections) # Set index to one past the last section
        st.session_state.ai_status = "All sections are complete! Proceed to Step 3 & 4."


# App layout
st.title("📄 Project Management Plan Generator")
st.markdown("Use this interactive tool to draft your Project Management Plan **one section at a time** using AI-generated suggestions, then refine and export.")

st.header("Step 1: Enter Project Description")
st.session_state.section_2_1 = st.text_area(
    "Describe your project (e.g., goals, high-level scope, key stakeholders).", 
    st.session_state.section_2_1,
    height=150
)

# --- Sequential Generation and Refinement (Step 2) ---
if st.session_state.section_2_1 and st.session_state.current_section_index < len(plan_sections):
    
    current_section = plan_sections[st.session_state.current_section_index]
    
    st.header(f"Step 2: Refine Section {st.session_state.current_section_index + 1}/{len(plan_sections)}: {current_section}")
    
    if not st.session_state.current_section_generated:
        st.button(f"Generate AI Suggestion for {current_section}", on_click=run_current_ai_suggestion)
    
    if st.session_state.ai_status:
        st.info(st.session_state.ai_status)
    
    # Display editable text area for the current section
    # Use a unique key to prevent content confusion on index change
    edited_text = st.text_area(
        f"Content for {current_section}", 
        st.session_state.sections.get(current_section, ""), 
        key=f"editor_{current_section}",
        height=300,
        disabled=(not st.session_state.current_section_generated) # Disable editing until content is generated
    )
    st.session_state.sections[current_section] = edited_text
    
    if st.session_state.current_section_generated:
        
        # Checkbox to accept the content
        accepted = st.checkbox(
            f"Accept Content for **{current_section}** (Ready to move to next section)",
            key=f"accept_{current_section}"
        )
        
        if accepted:
            st.button(
                f"✅ Accept & Move to Next Section", 
                on_click=accept_and_move_next,
                type="primary"
            )

# --- Final Steps (Step 3 & 4) ---
if st.session_state.current_section_index >= len(plan_sections):
    
    st.header("Step 3: Preview Final Document")
    with st.expander("Click to view full, completed plan preview"):
        st.markdown(f"## Project Management Plan")
        st.markdown(f"**Project Description:**\n{st.session_state.section_2_1}")
        for section in plan_sections:
            st.markdown(f"### {section}")
            st.markdown(st.session_state.sections[section])

    st.header("Step 4: Export Document")

    # --- Export to Word (.docx) ---
    if st.button("Export to Word (.docx)"):
        doc = Document()
        doc.add_heading("Project Management Plan", level=1)
        
        # Corrected f-string usage
        doc.add_paragraph(f"Project Description:\n{st.session_state.section_2_1}") 
        
        for section in plan_sections:
            doc.add_heading(section, level=2)
            doc.add_paragraph(st.session_state.sections[section])
        
        # Save document to an in-memory buffer
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        
        st.download_button(
            "Download Word Document", 
            word_buffer, 
            file_name="Project_Management_Plan.docx"
        )

    # --- Export to PDF (.pdf) ---
    if st.button("Export to PDF (.pdf)"):
        pdf_buffer = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        
        # Compile text content
        text = f"Project Management Plan\n\nProject Description:\n{st.session_state.section_2_1}\n\n"
        for section in plan_sections:
            text += f"{section}\n\n{st.session_state.sections[section]}\n\n"
        
        # Insert text onto the PDF page
        page.insert_text((72, 72), text, fontsize=12)
        doc.save(pdf_buffer)
        doc.close() # Close the document after saving
        pdf_buffer.seek(0)
        
        st.download_button(
            "Download PDF Document", 
            pdf_buffer, 
            file_name="Project_Management_Plan.pdf"
        )
