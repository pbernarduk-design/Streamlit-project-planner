import streamlit as st
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import io

# Initialize session state
if "section_2_1" not in st.session_state:
    st.session_state.section_2_1 = ""
if "sections" not in st.session_state:
    st.session_state.sections = {}

# Define sections of the project management plan
plan_sections = [
    "Executive Summary",
    "Outcomes & Vision",
    "Objectives",
    "Scope & Deliverables",
    "Governance & Team",
    "Risk Management"
]

# Function to generate AI suggestion (placeholder logic)
def generate_ai_suggestion(section_title, description):
    return f"Suggested content for {section_title} based on project description: '{description}'."

# App layout
st.title("📄 Project Management Plan Generator")

st.header("Step 1: Enter Project Description (Section 2.1)")
st.session_state.section_2_1 = st.text_area("Project Description", st.session_state.section_2_1)

if st.session_state.section_2_1:
    st.header("Step 2: Review and Edit Each Section")
    for section in plan_sections:
        st.subheader(section)
        if section not in st.session_state.sections:
            st.session_state.sections[section] = generate_ai_suggestion(section, st.session_state.section_2_1)
        edited_text = st.text_area(f"Edit {section}", st.session_state.sections[section], key=section)
        st.session_state.sections[section] = edited_text

    st.header("Step 3: Preview Final Document")
    for section in plan_sections:
        st.markdown(f"### {section}")
        st.markdown(st.session_state.sections[section])

    st.header("Step 4: Export Document")

    # Export to Word
    if st.button("Export to Word"):
        doc = Document()
        doc.add_picture("logo.png", width=Inches(1.5))
        doc.add_heading("Project Management Plan", level=1)
        doc.add_paragraph(f"Project Description:
{st.session_state.section_2_1}")
        for section in plan_sections:
            doc.add_heading(section, level=2)
            doc.add_paragraph(st.session_state.sections[section])
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button("Download Word Document", word_buffer, file_name="Project_Management_Plan.docx")

    # Export to PDF
    if st.button("Export to PDF"):
        pdf_buffer = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        text = f"Project Management Plan\n\nProject Description:\n{st.session_state.section_2_1}\n\n"
        for section in plan_sections:
            text += f"{section}\n{st.session_state.sections[section]}\n\n"
        page.insert_text((72, 72), text, fontsize=12)
        doc.save(pdf_buffer)
        pdf_buffer.seek(0)
        st.download_button("Download PDF Document", pdf_buffer, file_name="Project_Management_Plan.pdf")
